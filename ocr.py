import markdownify
import numpy as np
import requests
from paddleocr import PaddleOCR, PPStructure
import os
from docx import Document
import mammoth
import subprocess
import fitz
import logging

from paddleocr.ppstructure.recovery.recovery_to_doc import sorted_layout_boxes
from pdf2image import convert_from_path

logging.basicConfig(format='[%(asctime)s %(filename)s:%(lineno)d] %(levelname)s: %(message)s', level=logging.INFO, force=True)

angle_cls_url = 'https://api-sit.piaozone.com/nlp_service/match/file'

USING_STRUCTURE = False

# 初始化一个OCR或structure识别器,todo: 这里为省去启动时间，后续考虑并发
if USING_STRUCTURE:
    pp_structure = PPStructure(show_log=True)
else:
    paddle_ocr = PaddleOCR(use_angle_cls=True, use_gpu=False, cpu_threads=8)


# 使pdf图片摆正
def up_image(pdf_path, pdf_document):
    files = {'file': open(pdf_path, 'rb')}
    # 通过http api识别角度
    response = requests.post(angle_cls_url, files=files)
    files['file'].close()

    response_data = response.json()
    if response_data['errcode'] != "0000":
        return

    # 读取 angle 的值
    angle = int(response_data['data']['angle'])
    logging.info(f"图像{pdf_path}倾斜度是:  {angle}")
    if angle < 30 or angle > 330:
        return

    # 默认只有一页
    pdf_document[0].set_rotation(360 - angle)
    pdf_document.save(pdf_path)
    logging.info(f"已摆正图像右旋转[{360 - angle}]度")


def before_ocr(file_path, filename, doc, page_no):
    # 单独提取第page_no页
    file_base, file_extension = os.path.splitext(filename)
    dest_filename = f"{file_base}_{page_no}{file_extension}"
    dest_path = os.path.join(file_path, 'tmp', dest_filename)
    if not os.path.exists(os.path.dirname(dest_path)):
        os.makedirs(os.path.dirname(dest_path))

    new_doc = fitz.open()
    new_doc.insert_pdf(doc, from_page=page_no-1, to_page=page_no-1)
    new_doc.save(dest_path)
    logging.info(f"已保存第{page_no}页到：{dest_path}")

    # 使pdf图片摆正
    up_image(dest_path, new_doc)

    new_doc.close()
    return dest_path


def structure_2_markdown(res):
    md_text = ""
    for region in res:
        if region['type'].lower() == 'table' and len(region['res']) > 0 and 'html' in region['res']:
            table_text = markdownify.markdownify(region['res']['html'])
            md_text += table_text
        elif region['type'].lower() == 'table_caption' and len(region['res']) > 0:
            for r in region['res']:
                if 'text' in r:
                    md_text = f"{md_text}## {r['text']} \n"

    return md_text


def pdf_structure(pdf_path):
    # Convert all pages in the PDF to images
    images = convert_from_path(pdf_path)
    img_cv = np.array(images[0])
    # Convert RGB to BGR
    img = img_cv[:, :, ::-1].copy()

    result = pp_structure(img)

    # 排序并转成Markdown
    h, w, _ = img.shape
    res = sorted_layout_boxes(result, w)
    md_text = structure_2_markdown(res)

    logging.info(f"已经提取{pdf_path}的structure信息：\n{md_text}")
    return md_text


# todo: ocr. 先写到一个pdf，在读(待优化）
def ocr(file_path, filename, doc, page_no):
    # 预处理
    dest_path = before_ocr(file_path, filename, doc, page_no)

    if USING_STRUCTURE:  # structure方式解析，但目前效果一般
        text = pdf_structure(dest_path)
    else:
        text = pdf_ocr(dest_path)

    return text


def pdf_ocr(img_path):
    global paddle_ocr
    result = paddle_ocr.ocr(img_path)

    text_list = []
    for idx in range(len(result)):
        res = result[idx]
        for line in res:
            text_list.append(line[-1][0])

    ret = '\n'.join(text_list)

    return ret


def direct_docx(file_path):
    doc = Document(file_path)

    text_list = []
    for para in doc.paragraphs:
        text_list.append(para.text)

    for t in doc.tables:
        for r in t.rows:
            row_cnt = [cell.text for cell in r.cells]
            text_list.append(', '.join(row_cnt))

    ret = '\n'.join(text_list)
    print(ret)


# 定义一个忽略图片处理的函数
def ignore(image):
    return {"src": ""}  # 返回一个空的 src 属性，这样图片不会显示


def convert_docx_to_html(file_path):
    with open(file_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file, convert_image=mammoth.images.img_element(ignore))
        html_content = result.value  # 生成的 HTML 内容
        messages = result.messages  # 转换过程中的任何消息
        return html_content


def pdf_structure_cmd(path):
    det_folder = os.path.split(path)[0]
    # cmd = f"paddleocr --image_dir={path} --output={det_folder} --type=structure --recovery=true --use_pdf2docx_api=true"
    # print(cmd)
    # os.system(cmd)
    cmd = [
        "paddleocr",
        "--image_dir", path,
        "--output", det_folder,
        "--type", "structure",
        "--recovery", "true",
        "--use_pdf2docx_api", "true"
    ]

    subprocess.run(cmd)

    det_file = os.path.basename(path).split('.')[0] + '.docx'
    det_path = f"{det_folder}/{det_file}"

    return convert_docx_to_html(det_path)
