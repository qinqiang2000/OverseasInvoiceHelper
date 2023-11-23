import os
from docx import Document
from paddleocr import PPStructure,draw_structure_result,save_structure_res
from pdf2docx.converter import Converter
import pdf2docx
import fitz

path = "../uploads/Trend-tek Corporation Limited/TREND TEK INV23-7-11-1.pdf"
# path = "/Users/qinqiang02/workspace/python/OverseasInvoiceHelper/uploads/tmp/WT/WT单据6_1.pdf"
path = "/Users/qinqiang02/workspace/python/OverseasInvoiceHelper/uploads/WT TECHNOLOGY PTE. LTD/WT 91985260.pdf"


def test_pdf2docx(path):
    cv = Converter(path)
    cv.convert(ignore_page_error=False)
    cv.close()

test_pdf2docx(path)

def test_fitz(filename):
    doc = fitz.open(filename)
    page = doc.load_page(0)
    print(page.get_text("text"))


def pdf_structure(path):
    det_folder = os.path.split(path)[0]
    cmd = f"paddleocr --image_dir={path} --output={det_folder} --type=structure --recovery=true --use_pdf2docx_api=true"
    print(cmd)
    os.system(cmd)

    det_file = os.path.basename(path).split('.')[0] + '.docx'
    doc = Document(f"{det_folder}/{det_file}")

    text_list = []
    for para in doc.paragraphs:
        text_list.append(para.text)

    for t in doc.tables:
        for r in t.rows:
            row_cnt = [cell.text for cell in r.cells]
            text_list.append(', '.join(row_cnt))

    ret = '\n'.join(text_list)
    print(ret)

    return ret


# pdf_structure(path)

import mammoth


# 定义一个忽略图片处理的函数
def ignore(image):
    return {"src": ""}  # 返回一个空的 src 属性，这样图片不会显示


def convert_docx_to_html(file_path):
    with open(file_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file, convert_image=mammoth.images.img_element(ignore))
        html_content = result.value  # 生成的 HTML 内容
        messages = result.messages  # 转换过程中的任何消息
        return html_content


def convert_docx_to_markdown(file_path):
    with open(file_path, "rb") as docx_file:
        # 定义转换选项
        options = {
            "style_map": "p[style-name='Caption'] => div.caption",
            "convert_image": mammoth.images.img_element(ignore)
        }

        # 使用自定义选项进行转换
        result = mammoth.convert_to_markdown(docx_file, convert_image=mammoth.images.img_element(ignore))
        markdown = result.value  # 生成的 Markdown 文本
        messages = result.messages  # 转换过程中的任何消息
        return markdown


# 示例使用
# test_path = "../uploads/tmp/晨泰/INV230823-1_1.docx"
det_folder = os.path.split(path)[0]
det_file = os.path.basename(path).split('.')[0] + '.docx'
det_path = f"{det_folder}/{det_file}"
# new_text1 = convert_docx_to_html("../uploads/Trend-tek Corporation Limited/TREND TEK INV23-7-11-1.docx")
# new_text = convert_docx_to_markdown(det_path)
# print(new_text)

# with open(det_path, "rb") as docx_file:
#     result = mammoth.extract_raw_text(docx_file)
#     text = result.value # The raw text
#     print(text)
#     messages = result.messages # Any messages